# test_installation.py
"""
Test script for Legal Assistant Phase 1 installation
اسکریپت تست راه‌اندازی فاز ۱ دستیار حقوقی
"""

import sys
from pathlib import Path
import json
import traceback

# Add project root to Python path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

def test_imports():
    """Test if all required modules can be imported"""
    print("🔍 بررسی import modules...")
    
    try:
        from src.core.config import validate_config, get_config
        print("✅ src.core.config")
        
        from src.core.models import LegalDocument, LegalArticle, TextChunk
        print("✅ src.core.models")
        
        from src.utils.text_utils import PersianTextProcessor
        print("✅ src.utils.text_utils")
        
        from src.data_processing.document_splitter import DocumentSplitter
        print("✅ src.data_processing.document_splitter")
        
        from src.data_processing.document_parser import LegalDocumentParser
        print("✅ src.data_processing.document_parser")
        
        from src.data_processing.text_processor import AdvancedTextProcessor
        print("✅ src.data_processing.text_processor")
        
        from src.data_processing.chunker import IntelligentChunker
        print("✅ src.data_processing.chunker")
        
        from src.data_processing.metadata_generator import MetadataGenerator
        print("✅ src.data_processing.metadata_generator")
        
        return True
    
    except ImportError as e:
        print(f"❌ خطا در import: {str(e)}")
        print("\n🔧 راه‌حل‌های احتمالی:")
        print("   1. بررسی کنید که همه فایل‌ها در مسیر صحیح قرار دارند")
        print("   2. اجرای: pip install python-docx pydantic hazm tqdm")
        print("   3. بررسی syntax فایل‌ها")
        return False

def test_python_packages():
    """Test if required Python packages are installed"""
    print("\n🔍 بررسی پکیج‌های Python...")
    
    required_packages = [
        ('docx', 'python-docx'),
        ('pydantic', 'pydantic'),
        ('hazm', 'hazm'),
        ('tqdm', 'tqdm'),
        ('re', 'regex (built-in)'),
        ('json', 'json (built-in)'),
        ('pathlib', 'pathlib (built-in)'),
        ('datetime', 'datetime (built-in)')
    ]
    
    missing_packages = []
    
    for package, install_name in required_packages:
        try:
            __import__(package)
            print(f"✅ {package}")
        except ImportError:
            print(f"❌ {package} - نصب کنید: pip install {install_name}")
            missing_packages.append(install_name)
    
    if missing_packages:
        print(f"\n🔧 دستور نصب پکیج‌های مفقود:")
        print(f"pip install {' '.join([p for p in missing_packages if 'built-in' not in p])}")
        return False
    
    return True

def test_persian_processing():
    """Test Persian text processing capabilities"""
    print("\n🔍 بررسی پردازش متن فارسی...")
    
    try:
        from src.utils.text_utils import PersianTextProcessor
        
        processor = PersianTextProcessor()
        
        # Test text cleaning
        test_text = "   قانون    مقررات انتظامی هیئت علمی   ( مصوب 22/12/1364 )   "
        cleaned = processor.clean_text(test_text)
        print(f"✅ تمیزکاری متن: '{test_text[:30]}...' → '{cleaned[:30]}...'")
        
        # Test keyword extraction
        sample_content = "در این قانون اعضای هیئت علمی دانشگاه‌ها موظف به رعایت مقررات انتظامی هستند."
        keywords = processor.extract_keywords(sample_content, max_keywords=5)
        print(f"✅ استخراج کلیدواژه: {keywords}")
        
        # Test Persian validation
        is_valid = processor.is_valid_persian_text(sample_content)
        print(f"✅ اعتبارسنجی فارسی: {is_valid}")
        
        # Test normalization
        test_chars = "ك ي ء"
        normalized = processor.normalize_persian_text(test_chars)
        print(f"✅ نرمال‌سازی کاراکتر: '{test_chars}' → '{normalized}'")
        
        return True
    
    except Exception as e:
        print(f"❌ خطا در پردازش فارسی: {str(e)}")
        return False

def test_configuration():
    """Test system configuration"""
    print("\n🔍 بررسی تنظیمات سیستم...")
    
    try:
        from src.core.config import validate_config, get_config
        
        # Test config validation
        if validate_config():
            print("✅ تنظیمات سیستم معتبر است")
        else:
            print("⚠️ تنظیمات سیستم نیاز به بررسی دارد")
        
        # Test config loading
        config = get_config()
        print(f"✅ بارگذاری تنظیمات: {len(config)} بخش")
        
        # Test specific configs
        doc_config = config.get('document', {})
        if doc_config:
            print(f"✅ تنظیمات سند: {len(doc_config)} پارامتر")
        
        return True
    
    except Exception as e:
        print(f"❌ خطا در تنظیمات: {str(e)}")
        return False

def test_models():
    """Test Pydantic models"""
    print("\n🔍 بررسی مدل‌های داده...")
    
    try:
        from src.core.models import LegalDocument, LegalArticle, TextChunk, DocumentType, ApprovalAuthority
        
        # Test LegalArticle creation
        article = LegalArticle(
            number="ماده ۱",
            title="تعاریف",
            content="در این قانون کلمات و عبارات زیر در معانی مشروح مقابل آنها به کار می‌روند."
        )
        print(f"✅ مدل LegalArticle: تعداد کلمات = {article.word_count}")
        
        # Test LegalDocument creation
        document = LegalDocument(
            id="test_001",
            title="قانون نمونه",
            approval_date="01/01/1400",
            approval_authority="مجلس شورای اسلامی",
            document_type="قانون",
            standalone_articles=[article]
        )
        print(f"✅ مدل LegalDocument: {document.total_articles} ماده")
        
        # Test JSON serialization
        doc_json = document.json(ensure_ascii=False)
        json_obj = json.loads(doc_json)
        print(f"✅ سریالایزیشن JSON: {len(doc_json)} کاراکتر")
        
        # Test TextChunk
        chunk = TextChunk(
            id="test_chunk_001",
            document_id="test_001",
            content="این یک chunk تست است.",
            chunk_type="article",
            position=0
        )
        print(f"✅ مدل TextChunk: {chunk.character_count} کاراکتر")
        
        return True
    
    except Exception as e:
        print(f"❌ خطا در مدل‌ها: {str(e)}")
        traceback.print_exc()
        return False

def test_file_structure():
    """Test project file structure"""
    print("\n🔍 بررسی ساختار فایل‌ها...")
    
    required_dirs = [
        "data/raw",
        "data/processed", 
        "data/sample",
        "src/core",
        "src/data_processing",
        "src/utils",
        "scripts"
    ]
    
    required_files = [
        "src/__init__.py",
        "src/core/__init__.py",
        "src/core/config.py",
        "src/core/models.py",
        "src/utils/__init__.py",
        "src/utils/text_utils.py",
        "src/data_processing/__init__.py",
        "src/data_processing/document_splitter.py",
        "src/data_processing/document_parser.py",
        "src/data_processing/text_processor.py",
        "src/data_processing/chunker.py",
        "src/data_processing/metadata_generator.py",
        "scripts/process_documents.py"
    ]
    
    all_good = True
    
    # Check directories
    print("📁 بررسی پوشه‌ها:")
    for dir_path in required_dirs:
        if Path(dir_path).exists():
            print(f"   ✅ {dir_path}")
        else:
            print(f"   ❌ پوشه ناموجود: {dir_path}")
            all_good = False
    
    # Check files
    print("\n📄 بررسی فایل‌ها:")
    for file_path in required_files:
        if Path(file_path).exists():
            file_size = Path(file_path).stat().st_size
            print(f"   ✅ {file_path} ({file_size} bytes)")
        else:
            print(f"   ❌ فایل ناموجود: {file_path}")
            all_good = False
    
    return all_good

def test_sample_processing():
    """Test sample document processing"""
    print("\n🔍 تست پردازش نمونه...")
    
    try:
        from src.data_processing.document_splitter import DocumentSplitter
        from src.data_processing.text_processor import AdvancedTextProcessor
        from src.utils.text_utils import PersianTextProcessor
        
        # Test document splitter initialization
        splitter = DocumentSplitter()
        print("✅ DocumentSplitter آماده است")
        
        # Test text processor
        text_processor = AdvancedTextProcessor()
        
        # Test sample text processing
        sample_law = {
            'id': 'test_001',
            'title': 'قانون نمونه تست',
            'raw_content': 'ماده ۱ - این متن نمونه است برای تست پردازش.',
            'approval_date': '01/01/1400',
            'approval_authority': 'مجلس شورای اسلامی'
        }
        
        processed_law = text_processor.process_document_from_dict(sample_law.copy())
        print("✅ پردازش متن نمونه موفقیت‌آمیز")
        
        # Test Persian processor
        persian_processor = PersianTextProcessor()
        test_text = "این یک متن فارسی است."
        cleaned = persian_processor.clean_text(test_text)
        print(f"✅ پردازش فارسی: '{test_text}' → '{cleaned}'")
        
        return True
    
    except Exception as e:
        print(f"❌ خطا در تست پردازش: {str(e)}")
        traceback.print_exc()
        return False

def check_input_file():
    """Check if input file exists"""
    print("\n🔍 بررسی فایل ورودی...")
    
    input_file = Path("data/raw/Part2_Legals.docx")
    
    if input_file.exists():
        file_size = input_file.stat().st_size / (1024 * 1024)  # MB
        print(f"✅ فایل ورودی موجود است: {file_size:.1f} MB")
        
        # Test if file is readable
        try:
            from docx import Document
            doc = Document(str(input_file))
            para_count = len(doc.paragraphs)
            print(f"✅ فایل قابل خواندن است: {para_count} پاراگراف")
            return True
        except Exception as e:
            print(f"⚠️ فایل موجود است اما قابل خواندن نیست: {str(e)}")
            return False
    else:
        print("⚠️ فایل Part2_Legals.docx در data/raw/ پیدا نشد")
        print("   لطفاً فایل را در مسیر صحیح قرار دهید")
        print("   راهنما: راهنمای بارگذاری فایل‌ها را مطالعه کنید")
        return False

def test_write_permissions():
    """Test write permissions in project directories"""
    print("\n🔍 بررسی مجوزهای نوشتن...")
    
    test_dirs = ["data/processed", "data/sample"]
    all_good = True
    
    for test_dir in test_dirs:
        try:
            Path(test_dir).mkdir(exist_ok=True)
            test_file = Path(test_dir) / "write_test.tmp"
            test_file.write_text("test")
            test_file.unlink()
            print(f"✅ مجوز نوشتن در {test_dir}")
        except Exception as e:
            print(f"❌ مشکل مجوز نوشتن در {test_dir}: {str(e)}")
            all_good = False
    
    return all_good

def generate_test_report():
    """Generate a comprehensive test report"""
    print("\n" + "="*60)
    print("📋 گزارش نهایی تست راه‌اندازی")
    print("="*60)
    
    tests = [
        ("Python Packages", test_python_packages),
        ("Import Modules", test_imports),
        ("Persian Processing", test_persian_processing),
        ("Configuration", test_configuration),
        ("Data Models", test_models),
        ("File Structure", test_file_structure),
        ("Sample Processing", test_sample_processing),
        ("Write Permissions", test_write_permissions),
        ("Input File", check_input_file)
    ]
    
    results = {}
    total_tests = len(tests)
    passed_tests = 0
    
    for test_name, test_func in tests:
        print(f"\n--- {test_name} ---")
        try:
            result = test_func()
            results[test_name] = "PASS" if result else "FAIL"
            if result:
                passed_tests += 1
        except Exception as e:
            results[test_name] = f"ERROR: {str(e)}"
            print(f"❌ خطای غیرمنتظره در {test_name}: {str(e)}")
    
    # Print summary
    print(f"\n{'='*60}")
    print(f"📊 خلاصه نتایج:")
    print(f"   ✅ موفق: {passed_tests}/{total_tests}")
    print(f"   ❌ ناموفق: {total_tests - passed_tests}/{total_tests}")
    print(f"   📈 نرخ موفقیت: {(passed_tests/total_tests)*100:.1f}%")
    
    # Detailed results
    print(f"\n📋 نتایج تفصیلی:")
    for test_name, result in results.items():
        status_icon = "✅" if result == "PASS" else "❌"
        print(f"   {status_icon} {test_name}: {result}")
    
    # Recommendations
    print(f"\n💡 توصیه‌ها:")
    if passed_tests == total_tests:
        print("   🎉 همه تست‌ها موفقیت‌آمیز! آماده اجرای فاز ۱ هستید.")
        print("   ▶️  برای شروع پردازش:")
        print("       python scripts/process_documents.py data/raw/Part2_Legals.docx")
        print("   ▶️  برای نمایش قابلیت‌ها:")
        print("       python demo_phase1.py")
    else:
        failed_tests = [name for name, result in results.items() if result != "PASS"]
        
        if "Python Packages" in failed_tests:
            print("   📦 ابتدا پکیج‌های مورد نیاز را نصب کنید:")
            print("       pip install python-docx pydantic hazm tqdm regex")
        
        if "File Structure" in failed_tests:
            print("   📂 ساختار پوشه‌ها را کامل کنید:")
            print("       mkdir -p data/{raw,processed,sample} src/{core,data_processing,utils} scripts")
        
        if "Input File" in failed_tests:
            print("   📁 فایل Part2_Legals.docx را در data/raw/ قرار دهید")
        
        if "Import Modules" in failed_tests:
            print("   🔧 فایل‌های کد را از artifacts کپی کنید")
        
        if "Write Permissions" in failed_tests:
            print("   🔐 مجوزهای نوشتن را بررسی کنید:")
            print("       chmod 755 data/")
    
    return results

def main():
    """Main function"""
    print("🚀 شروع تست راه‌اندازی دستیار حقوقی هوشمند - فاز ۱")
    print("="*60)
    print(f"📍 مسیر پروژه: {Path.cwd()}")
    print(f"🐍 نسخه Python: {sys.version}")
    
    # Generate comprehensive test report
    results = generate_test_report()
    
    # Save test results
    try:
        test_results_file = Path("test_results.json")
        with open(test_results_file, 'w', encoding='utf-8') as f:
            json.dump({
                'timestamp': str(datetime.now().isoformat() if 'datetime' in globals() else 'unknown'),
                'python_version': sys.version,
                'project_path': str(Path.cwd()),
                'results': results,
                'total_tests': len(results),
                'passed_tests': sum(1 for r in results.values() if r == "PASS"),
                'success_rate': (sum(1 for r in results.values() if r == "PASS") / len(results)) * 100
            }, f, ensure_ascii=False, indent=2)
        
        print(f"\n💾 نتایج تست در فایل ذخیره شد: {test_results_file}")
    
    except Exception as e:
        print(f"\n⚠️ خطا در ذخیره نتایج: {str(e)}")
    
    print(f"\n{'='*60}")
    print("✨ تست راه‌اندازی تکمیل شد!")

if __name__ == "__main__":
    try:
        from datetime import datetime
        main()
    except KeyboardInterrupt:
        print("\n\n⏹️ تست توسط کاربر متوقف شد")
    except Exception as e:
        print(f"\n💥 خطای غیرمنتظره: {str(e)}")
        traceback.print_exc()